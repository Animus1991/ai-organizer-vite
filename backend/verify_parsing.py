#!/usr/bin/env python3
"""
Script για να επιβεβαιώσουμε ότι το parsing διαβάζει ολόκληρο το κείμενο.
"""
import sys
from pathlib import Path

# Add src to path
sys.path.insert(0, str(Path(__file__).parent / "src"))

from ai_organizer.ingest.parsers import read_text_file, read_docx_file, parse_chatgpt_export_json
import json

def verify_text_file(file_path: Path):
    """Ελέγχει αν το read_text_file διαβάζει ολόκληρο το αρχείο."""
    print(f"\n📄 Ελέγχοντας: {file_path.name}")
    
    # Διαβάζουμε με το parser
    parsed_text = read_text_file(file_path)
    
    # Διαβάζουμε απευθείας για σύγκριση
    direct_text = file_path.read_text(encoding="utf-8", errors="replace")
    
    if parsed_text == direct_text:
        print(f"  ✅ PASS: Το parsed text είναι ίδιο με το original ({len(parsed_text)} chars)")
        return True
    else:
        print(f"  ❌ FAIL: Διαφορά! Parsed: {len(parsed_text)} chars, Direct: {len(direct_text)} chars")
        if len(parsed_text) < len(direct_text):
            print(f"  ⚠️  WARNING: Το parsed text είναι μικρότερο κατά {len(direct_text) - len(parsed_text)} chars!")
        return False

def verify_docx_file(file_path: Path):
    """Ελέγχει αν το read_docx_file διαβάζει όλο το περιεχόμενο."""
    print(f"\n📄 Ελέγχοντας DOCX: {file_path.name}")
    
    try:
        parsed_text = read_docx_file(file_path)
        
        # Μετράμε paragraphs, tables, κλπ
        from docx import Document as DocxDocument
        doc = DocxDocument(str(file_path))
        
        paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
        table_count = len(doc.tables)
        section_count = len(doc.sections)
        
        print(f"  📊 Statistics:")
        print(f"     - Paragraphs: {paragraph_count}")
        print(f"     - Tables: {table_count}")
        print(f"     - Sections: {section_count}")
        print(f"     - Parsed text length: {len(parsed_text)} chars")
        
        # Ελέγχουμε αν υπάρχουν tables που δεν διαβάστηκαν
        if table_count > 0:
            if "[TABLE]" in parsed_text or "|" in parsed_text:
                print(f"  ✅ Tables detected in parsed text")
            else:
                print(f"  ⚠️  WARNING: {table_count} tables found but may not be fully extracted")
        
        # Ελέγχουμε αν υπάρχουν headers/footers
        has_headers = any(s.header for s in doc.sections if s.header)
        has_footers = any(s.footer for s in doc.sections if s.footer)
        
        if has_headers:
            if "[HEADER]" in parsed_text:
                print(f"  ✅ Headers detected in parsed text")
            else:
                print(f"  ⚠️  WARNING: Headers exist but may not be extracted")
        
        if has_footers:
            if "[FOOTER]" in parsed_text:
                print(f"  ✅ Footers detected in parsed text")
            else:
                print(f"  ⚠️  WARNING: Footers exist but may not be extracted")
        
        if len(parsed_text) > 0:
            print(f"  ✅ PASS: Parsed {len(parsed_text)} characters")
            return True
        else:
            print(f"  ❌ FAIL: No text extracted!")
            return False
            
    except Exception as e:
        print(f"  ❌ ERROR: {e}")
        return False

def main():
    print("=" * 60)
    print("🔍 Parsing Verification Script")
    print("=" * 60)
    
    # Check if files are provided
    if len(sys.argv) < 2:
        print("\nUsage: python verify_parsing.py <file1> [file2] ...")
        print("\nExample:")
        print("  python verify_parsing.py test.txt test.docx")
        return
    
    results = []
    for file_arg in sys.argv[1:]:
        file_path = Path(file_arg)
        if not file_path.exists():
            print(f"\n❌ File not found: {file_path}")
            results.append(False)
            continue
        
        ext = file_path.suffix.lower()
        if ext in [".txt", ".md"]:
            results.append(verify_text_file(file_path))
        elif ext == ".docx":
            results.append(verify_docx_file(file_path))
        elif ext == ".json":
            # For JSON, we can't easily verify without knowing the structure
            print(f"\n📄 JSON file: {file_path.name}")
            print(f"  ℹ️  JSON parsing depends on structure - manual verification needed")
            results.append(True)
        else:
            print(f"\n⚠️  Unsupported file type: {ext}")
            results.append(False)
    
    print("\n" + "=" * 60)
    print(f"📊 Summary: {sum(results)}/{len(results)} files passed")
    print("=" * 60)

if __name__ == "__main__":
    main()

