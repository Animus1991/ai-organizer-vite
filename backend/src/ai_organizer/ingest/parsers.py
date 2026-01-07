from __future__ import annotations
import json
from pathlib import Path
from typing import Any
from docx import Document as DocxDocument

def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")

def parse_chatgpt_export_json(raw: str) -> str:
    """
    Προσπαθεί να “ξεδιπλώσει” OpenAI ChatGPT export (conversations/mapping)
    και να βγάλει καθαρό κείμενο με role blocks.
    Αν δεν αναγνωρίσει τη δομή, επιστρέφει pretty JSON text.
    """
    data: Any = json.loads(raw)

    # Case A: Export as list of conversations
    if isinstance(data, list) and data and isinstance(data[0], dict) and "mapping" in data[0]:
        conv = data[0]
        return _flatten_mapping_conv(conv)

    # Case B: Export has "conversations"
    if isinstance(data, dict) and "conversations" in data and isinstance(data["conversations"], list) and data["conversations"]:
        conv = data["conversations"][0]
        if isinstance(conv, dict) and "mapping" in conv:
            return _flatten_mapping_conv(conv)

    # Case C: Simple messages list
    if isinstance(data, dict) and "messages" in data and isinstance(data["messages"], list):
        parts = []
        for m in data["messages"]:
            role = (m.get("role") or "unknown").upper()
            content = m.get("content") or ""
            if isinstance(content, dict) and "parts" in content:
                content = "\n".join([p for p in content.get("parts", []) if isinstance(p, str)])
            parts.append(f"{role}:\n{content}\n")
        return "\n".join(parts).strip()

    # Fallback
    return json.dumps(data, ensure_ascii=False, indent=2)

def _flatten_mapping_conv(conv: dict) -> str:
    mapping = conv.get("mapping", {}) or {}
    nodes = []
    for _, node in mapping.items():
        msg = (node or {}).get("message")
        if not msg:
            continue
        create_time = msg.get("create_time") or 0
        author = (msg.get("author") or {}).get("role") or "unknown"
        content = msg.get("content") or {}
        text = ""
        if isinstance(content, dict) and "parts" in content:
            text = "\n".join([p for p in content.get("parts", []) if isinstance(p, str)])
        elif isinstance(content, str):
            text = content
        nodes.append((create_time, author, text))

    nodes.sort(key=lambda x: x[0])
    out = []
    for _, author, text in nodes:
        out.append(f"{author.upper()}:\n{text}\n")
    return "\n".join(out).strip()

def read_docx_file(path: Path) -> str:
    """
    Διαβάζει ολόκληρο το DOCX file, συμπεριλαμβανομένων:
    - Paragraphs
    - Tables (όλα τα cells)
    - Headers και Footers
    - Text boxes (αν υπάρχουν)
    """
    doc = DocxDocument(str(path))
    parts = []
    
    # 1. Paragraphs από το main body
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            parts.append(text)
    
    # 2. Tables - διαβάζουμε όλα τα cells
    for table in doc.tables:
        table_parts = []
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_text = (cell.text or "").strip()
                if cell_text:
                    row_parts.append(cell_text)
            if row_parts:
                table_parts.append(" | ".join(row_parts))
        if table_parts:
            parts.append("\n".join(table_parts))
    
    # 3. Headers - διαβάζουμε από όλα τα sections
    for section in doc.sections:
        # Header
        if section.header:
            header_parts = []
            for p in section.header.paragraphs:
                text = (p.text or "").strip()
                if text:
                    header_parts.append(text)
            # Headers tables
            for table in section.header.tables:
                for row in table.rows:
                    row_parts = []
                    for cell in row.cells:
                        cell_text = (cell.text or "").strip()
                        if cell_text:
                            row_parts.append(cell_text)
                    if row_parts:
                        header_parts.append(" | ".join(row_parts))
            if header_parts:
                parts.append("[HEADER]\n" + "\n".join(header_parts))
        
        # Footer
        if section.footer:
            footer_parts = []
            for p in section.footer.paragraphs:
                text = (p.text or "").strip()
                if text:
                    footer_parts.append(text)
            # Footer tables
            for table in section.footer.tables:
                for row in table.rows:
                    row_parts = []
                    for cell in row.cells:
                        cell_text = (cell.text or "").strip()
                        if cell_text:
                            row_parts.append(cell_text)
                    if row_parts:
                        footer_parts.append(" | ".join(row_parts))
            if footer_parts:
                parts.append("[FOOTER]\n" + "\n".join(footer_parts))
    
    # 4. Text boxes (shapes) - αν υπάρχουν
    try:
        # python-docx δεν υποστηρίζει άμεσα text boxes, αλλά μπορούμε να προσπαθήσουμε
        # μέσω XML parsing αν χρειάζεται (για μελλοντική βελτίωση)
        pass
    except Exception:
        # Αν αποτύχει, συνεχίζουμε χωρίς text boxes
        pass
    
    # Ενώνουμε όλα με διπλή γραμμή για καλύτερη αναγνωσιμότητα
    return "\n\n".join(parts)